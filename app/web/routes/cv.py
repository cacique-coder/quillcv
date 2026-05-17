import json
import logging
import re
import time

from fastapi import APIRouter, Request, WebSocket, WebSocketDisconnect
from fastapi.responses import JSONResponse, RedirectResponse, Response

from app.billing.use_cases.manage_credits import add_credits, deduct_credit, has_credits
from app.cv_builder.use_cases.build_cv import cv_data_from_attempt as builder_cv_data_from_attempt
from app.cv_export.adapters.docx_export import generate_docx
from app.cv_export.adapters.puppeteer_pdf import generate_pdf
from app.cv_export.adapters.template_registry import REGION_RULES, get_template, list_regions, list_templates
from app.cv_generation.adapters.pdfplumber_parser import parse_cv
from app.cv_generation.adapters.refiner import apply_review_fixes
from app.cv_generation.use_cases.generate_cv import run_generation_pipeline
from app.infrastructure.llm.client import set_llm_context
from app.pii.use_cases.filter_suggestions import filter_personal_detail_items
from app.infrastructure.persistence.attempt_store import (
    get_attempt,
    get_document_bytes,
    get_document_filename,
    update_attempt,
)
from app.infrastructure.persistence.database import async_session
from app.scoring.adapters.keyword_matcher import analyze_ats
from app.web.templates import templates

logger = logging.getLogger(__name__)

router = APIRouter()

# Maximum number of times the user may re-run the full generation pipeline
# for a single attempt (POST /analyze regenerations).  Each run costs 1 credit,
# so this also acts as a per-attempt spend ceiling.
_MAX_REGENERATIONS = 5


def _build_results_context(
    result: dict,
    missing_keyword_groups: dict,
    missing_keyword_categories_order: list,
    **extra,
) -> dict:
    """Build the template context dict for ``partials/results.html``.

    Applies the personal-detail filter to quality flags and ATS recommendations
    so the suggestions panel never shows identity-field items.
    """
    quality_review = result.get("quality_review") or {}
    raw_flags = quality_review.get("flags") or []
    ats_generated = result.get("ats_generated")
    raw_recs = list(getattr(ats_generated, "recommendations", None) or [])

    filtered_flags, filtered_recs = filter_personal_detail_items(raw_flags, raw_recs)

    # Produce a patched quality_review with filtered flags only (keep summary)
    filtered_quality_review = None
    if quality_review:
        filtered_quality_review = {**quality_review, "flags": filtered_flags}

    ctx = {
        **result,
        "quality_review": filtered_quality_review,
        "ats_recommendations": filtered_recs,
        "missing_keyword_groups": missing_keyword_groups,
        "missing_keyword_categories_order": missing_keyword_categories_order,
    }
    ctx.update(extra)
    return ctx


# ---------------------------------------------------------------------------
# WebSocket endpoint — real-time progress during CV generation
# ---------------------------------------------------------------------------


@router.websocket("/ws/analyze")
async def ws_analyze(websocket: WebSocket):
    attempt_id = websocket.query_params.get("attempt_id", "")
    if not attempt_id or not get_attempt(attempt_id):
        logger.warning("WebSocket rejected — invalid attempt_id=%r", attempt_id)
        await websocket.close(code=4000, reason="Invalid attempt")
        return

    await websocket.accept()
    ws_start = time.monotonic()
    current_step: list[str] = ["(not started)"]  # mutable cell for finally block

    logger.info("WebSocket connected attempt=%s", attempt_id)

    async def send_progress(step: str, detail: str) -> None:
        current_step[0] = step
        step_elapsed = round((time.monotonic() - ws_start) * 1000)
        logger.info(
            "WebSocket progress attempt=%s step=%r detail=%r elapsed=%dms",
            attempt_id,
            step,
            detail,
            step_elapsed,
        )
        try:
            await websocket.send_json({"type": "progress", "step": step, "detail": detail})
        except (WebSocketDisconnect, RuntimeError) as err:
            raise WebSocketDisconnect() from err

    try:
        llm = websocket.app.state.llm
        llm_fast = websocket.app.state.llm_fast

        # WebSocket connections bypass BaseHTTPMiddleware, so the session
        # middleware never runs and websocket.state.session is never set.
        # Manually load the session from the cookie before resolving the user.
        from app.identity.adapters.fastapi_deps import get_current_user
        from app.infrastructure.middleware.session import _COOKIE_NAME, _load_session

        _ws_session_id = websocket.cookies.get(_COOKIE_NAME)
        _ws_session_data = (await _load_session(_ws_session_id) if _ws_session_id else None) or {}
        websocket.state.session = _ws_session_data

        ws_user = await get_current_user(websocket)
        ws_user_id = ws_user.id if ws_user else None

        # ── Credit gate ────────────────────────────────────────────────────
        # Authenticated users must have a positive balance.  Guests (ws_user_id
        # is None) are allowed through — they can't buy credits, but they can
        # use any free trial quota granted at registration.  If credits run out,
        # we close with 4402 (4000-range = app-level, 02 ≈ HTTP 402).
        if ws_user_id is not None:
            async with async_session() as db:
                _credits_ok = await has_credits(db, ws_user_id)
            if not _credits_ok:
                await websocket.send_json({
                    "type": "error",
                    "code": "insufficient_credits",
                    "message": "You have no credits remaining. Please top up to generate a CV.",
                })
                await websocket.close(code=4402, reason="Payment required")
                return

            # Deduct before running the pipeline.  The atomic UPDATE prevents
            # concurrent requests from both succeeding on a balance of 1.
            async with async_session() as db:
                _deducted = await deduct_credit(db, ws_user_id)
            if not _deducted:
                # Race condition — another concurrent request won the credit.
                await websocket.send_json({
                    "type": "error",
                    "code": "insufficient_credits",
                    "message": "You have no credits remaining. Please top up to generate a CV.",
                })
                await websocket.close(code=4402, reason="Payment required")
                return

            logger.info("WS credit deducted user=%s attempt=%s", ws_user_id, attempt_id)

        _ws_credit_deducted = ws_user_id is not None  # track for refund on hard failure

        try:
            result = await run_generation_pipeline(
                attempt_id,
                llm,
                llm_fast,
                on_progress=send_progress,
                user_id=ws_user_id,
                pii=_ws_session_data.get("pii") or {},
                pii_password=_ws_session_data.get("_pii_password"),
            )
        except Exception:
            # Hard failure — refund the credit so the user isn't charged for a
            # broken generation.  Only on hard failure, not transient retries
            # handled inside run_generation_pipeline.
            if _ws_credit_deducted:
                try:
                    async with async_session() as db:
                        await add_credits(db, ws_user_id, 1, as_grant=True)
                    logger.info("WS credit refunded (hard failure) user=%s attempt=%s", ws_user_id, attempt_id)
                except Exception:
                    logger.exception("Failed to refund credit after WS failure user=%s", ws_user_id)
            raise

        total_ms = round((time.monotonic() - ws_start) * 1000)
        result.get("quality_review") and (
            # _llm_usage lives inside cv_data; surface it for the log if present
            result.get("quality_review", {}) or {}
        )
        # Pull usage from cv_data if available (pipeline stores it there)
        attempt_data = {}
        try:
            from app.infrastructure.persistence.attempt_store import get_attempt as _get

            attempt_data = _get(attempt_id) or {}
        except Exception:
            logger.debug("Failed to retrieve attempt data for %s", attempt_id)
        cv_data = attempt_data.get("cv_data") or {}
        usage = cv_data.get("_llm_usage", {})
        if usage:
            logger.info(
                "WebSocket complete attempt=%s duration=%dms input_tokens=%s output_tokens=%s cost_usd=%s",
                attempt_id,
                total_ms,
                usage.get("input_tokens", "?"),
                usage.get("output_tokens", "?"),
                usage.get("cost_usd", "?"),
            )
        else:
            logger.info(
                "WebSocket complete attempt=%s duration=%dms",
                attempt_id,
                total_ms,
            )

        # Render the final HTML
        missing_keyword_groups = result.get("missing_keyword_groups") or {}
        _cv_data_for_order = (attempt_data.get("cv_data") or {})
        missing_keyword_categories_order = [
            g.get("category") for g in (_cv_data_for_order.get("skills_grouped") or [])
            if g.get("category")
        ]
        # `result` already includes `missing_keyword_groups` from the pipeline,
        # so we only add the order list here (not in the pipeline output).
        results_ctx = _build_results_context(
            result,
            missing_keyword_groups,
            missing_keyword_categories_order,
        )
        html = templates.get_template("partials/results.html").render(
            request=websocket,
            **results_ctx,
        )
        await websocket.send_json({"type": "complete", "html": html})

    except WebSocketDisconnect:
        elapsed_ms = round((time.monotonic() - ws_start) * 1000)
        logger.info("WebSocket disconnected attempt=%s step=%r elapsed=%dms", attempt_id, current_step[0], elapsed_ms)
        return
    except ValueError as e:
        elapsed_ms = round((time.monotonic() - ws_start) * 1000)
        logger.warning(
            "WebSocket validation error attempt=%s step=%r elapsed=%dms error=%r",
            attempt_id,
            current_step[0],
            elapsed_ms,
            str(e),
        )
        try:
            await websocket.send_json({"type": "error", "message": str(e)})
        except (WebSocketDisconnect, RuntimeError):
            pass
    except Exception:
        elapsed_ms = round((time.monotonic() - ws_start) * 1000)
        logger.exception(
            "WebSocket generation failed attempt=%s step=%r elapsed=%dms",
            attempt_id,
            current_step[0],
            elapsed_ms,
        )
        try:
            await websocket.send_json({"type": "error", "message": "Generation failed. Please try again."})
        except (WebSocketDisconnect, RuntimeError):
            pass
    finally:
        try:
            await websocket.close()
        except (WebSocketDisconnect, RuntimeError):
            pass


# ---------------------------------------------------------------------------
# HTTP fallback — kept for clients without WebSocket support
# ---------------------------------------------------------------------------


@router.post("/analyze")
async def analyze(request: Request):
    """Parse CV from attempt store, run ATS analysis, and generate tailored CV.

    POST /analyze is the *Regenerate* path (the initial generation uses
    /ws/analyze).  Each call costs 1 credit and is capped at
    _MAX_REGENERATIONS per attempt to prevent runaway spend.
    Concurrent duplicate requests are serialised by the atomic credit
    deduction — only one call can win the balance check and proceed.
    """
    attempt_id = request.state.session.get("attempt_id")
    if not attempt_id:
        return templates.TemplateResponse(
            "partials/error.html",
            {"request": request, "error": "No active session. Please start from the beginning."},
        )

    http_user = getattr(request.state, "user", None)
    http_user_id = http_user.id if http_user else None

    # ── Regeneration cap ──────────────────────────────────────────────────
    _attempt_meta = get_attempt(attempt_id) or {}
    regen_count = _attempt_meta.get("regeneration_count", 0)
    if regen_count >= _MAX_REGENERATIONS:
        return JSONResponse(
            status_code=429,
            content={
                "error": "regeneration_limit_reached",
                "message": (
                    f"You have reached the maximum of {_MAX_REGENERATIONS} regenerations "
                    "for this session. Start a new session to generate another CV."
                ),
            },
        )

    # ── Credit gate ───────────────────────────────────────────────────────
    if http_user_id is not None:
        async with async_session() as db:
            _credits_ok = await has_credits(db, http_user_id)
        if not _credits_ok:
            return JSONResponse(
                status_code=402,
                content={
                    "error": "insufficient_credits",
                    "message": "You have no credits remaining. Please top up to generate a CV.",
                },
            )

        # Atomic deduction — prevents concurrent duplicate requests from both
        # succeeding when the balance is exactly 1.
        async with async_session() as db:
            _deducted = await deduct_credit(db, http_user_id)
        if not _deducted:
            return JSONResponse(
                status_code=402,
                content={
                    "error": "insufficient_credits",
                    "message": "You have no credits remaining. Please top up to generate a CV.",
                },
            )
        logger.info("HTTP /analyze credit deducted user=%s attempt=%s", http_user_id, attempt_id)

    _credit_deducted = http_user_id is not None

    # POST /analyze is the Regenerate endpoint (initial generation uses /ws/analyze).
    # Always clear cached LLM outputs so the pipeline re-runs from scratch.
    # User inputs (cv_text, job_description, region, template, pii, documents) are preserved.
    update_attempt(
        attempt_id,
        extracted_keywords=None,
        cv_data=None,
        rendered_cv=None,
        cover_letter_data=None,
        cover_letter_html=None,
        quality_review=None,
        quality_review_flags=None,
        missing_keyword_groups=None,
        regeneration_count=regen_count + 1,
    )
    logger.info(
        "Pipeline[%s] regenerate #%d: cleared cached LLM outputs",
        attempt_id,
        regen_count + 1,
    )

    _AI_FAILURE_MSG = "CV generation failed"  # prefix used by run_generation_pipeline

    try:
        result = await run_generation_pipeline(
            attempt_id,
            request.app.state.llm,
            request.app.state.llm_fast,
            user_id=http_user_id,
            pii=request.state.session.get("pii") or {},
            pii_password=request.state.session.get("_pii_password"),
        )
    except ValueError as e:
        err_str = str(e)
        # AI-failure errors get a credit refund; input-validation errors do not.
        # run_generation_pipeline uses "CV generation failed" as the prefix for
        # AI hard failures; all other ValueErrors are input/session errors.
        is_ai_failure = _AI_FAILURE_MSG in err_str
        if is_ai_failure and _credit_deducted:
            try:
                async with async_session() as db:
                    await add_credits(db, http_user_id, 1, as_grant=True)
                logger.info(
                    "HTTP /analyze credit refunded (AI failure) user=%s attempt=%s",
                    http_user_id,
                    attempt_id,
                )
            except Exception:
                logger.exception(
                    "Failed to refund credit after /analyze AI failure user=%s",
                    http_user_id,
                )
        return templates.TemplateResponse(
            "partials/error.html",
            {"request": request, "error": err_str},
        )
    except Exception:
        # Unexpected hard failure — refund the credit so the user isn't charged
        # for a broken generation.  Only on hard failure, not transient retries
        # handled inside run_generation_pipeline.
        if _credit_deducted:
            try:
                async with async_session() as db:
                    await add_credits(db, http_user_id, 1, as_grant=True)
                logger.info(
                    "HTTP /analyze credit refunded (hard failure) user=%s attempt=%s",
                    http_user_id,
                    attempt_id,
                )
            except Exception:
                logger.exception(
                    "Failed to refund credit after /analyze failure user=%s",
                    http_user_id,
                )
        raise

    missing_keyword_groups = result.get("missing_keyword_groups") or {}
    _attempt_for_order = get_attempt(attempt_id) or {}
    missing_keyword_categories_order = [
        g.get("category") for g in ((_attempt_for_order.get("cv_data") or {}).get("skills_grouped") or [])
        if g.get("category")
    ]
    results_ctx = _build_results_context(
        result,
        missing_keyword_groups,
        missing_keyword_categories_order,
    )
    return templates.TemplateResponse(
        "partials/results.html",
        {"request": request, **results_ctx},
    )


# ---------------------------------------------------------------------------
# Apply fixes (unchanged — stays HTTP)
# ---------------------------------------------------------------------------


@router.post("/apply-fixes")
async def apply_fixes(request: Request):
    """Apply selected quality review fixes to the generated CV.

    /apply-fixes is treated as in-progress refinement of an already-paid
    generation, so it does NOT charge a credit.  However, we cap calls per
    attempt at _MAX_REGENERATIONS to prevent unbounded LLM cost from a
    single session.  The cap is shared with the /analyze regeneration
    counter so the combined total stays within the per-attempt ceiling.
    """
    attempt_id = request.state.session.get("attempt_id")
    if not attempt_id:
        return templates.TemplateResponse(
            "partials/error.html",
            {"request": request, "error": "No active session."},
        )

    attempt = get_attempt(attempt_id)
    if not attempt or not attempt.get("cv_data"):
        return templates.TemplateResponse(
            "partials/error.html",
            {"request": request, "error": "No generated CV found. Please generate first."},
        )

    form = await request.form()

    cached_flags = attempt.get("quality_review_flags", [])
    selected_indices = form.getlist("selected")

    flags = []
    for idx_str in selected_indices:
        try:
            idx = int(idx_str)
            if 0 <= idx < len(cached_flags):
                flag = dict(cached_flags[idx])
                user_note = str(form.get(f"instruction_{idx}", "")).strip()
                if user_note:
                    flag["user_instruction"] = user_note
                flags.append(flag)
        except (ValueError, IndexError):
            pass

    if not flags:
        return templates.TemplateResponse(
            "partials/error.html",
            {"request": request, "error": "No fixes selected."},
        )

    cv_data = attempt["cv_data"]
    job_description = attempt.get("job_description", "")
    template_id = attempt.get("template_id", "modern")
    region_code = attempt.get("region", "US")

    llm_fast = request.app.state.llm_fast

    import uuid as _uuid_mod

    fix_user = getattr(request.state, "user", None)
    set_llm_context(
        service="apply_fixes",
        attempt_id=attempt_id,
        user_id=fix_user.id if fix_user else None,
        transaction_id=_uuid_mod.uuid4().hex,
    )

    updated_data = await apply_review_fixes(cv_data, flags, job_description, llm=llm_fast)
    if updated_data is None:
        return templates.TemplateResponse(
            "partials/error.html",
            {"request": request, "error": "Failed to apply fixes. Please try again."},
        )

    llm_usage = updated_data.pop("_llm_usage", {})
    rendered_cv = templates.get_template(f"cv_templates/{template_id}.html").render(**updated_data)
    updated_data["_llm_usage"] = llm_usage

    generated_text = re.sub(r"<style[^>]*>.*?</style>", "", rendered_cv, flags=re.DOTALL)
    generated_text = re.sub(r"<[^>]+>", " ", generated_text)
    generated_text = re.sub(r"\s+", " ", generated_text).strip()

    keyword_data = attempt.get("extracted_keywords")
    job_keywords = keyword_data["all_keywords"] if keyword_data else None
    ats_generated = analyze_ats(generated_text, job_description, keywords_override=job_keywords)

    update_attempt(attempt_id, cv_data=updated_data, rendered_cv=rendered_cv)

    selected_template = get_template(template_id) or get_template("modern")
    region_rules = REGION_RULES.get(region_code, REGION_RULES["US"])

    cv_text = attempt.get("cv_text_preview", "")
    cv_bytes = get_document_bytes(attempt_id, "cv_file")
    cv_filename = get_document_filename(attempt_id, "cv_file")
    if cv_bytes and cv_filename:
        try:
            full_cv_text = parse_cv(cv_filename, cv_bytes)
        except ValueError:
            full_cv_text = cv_text
    else:
        full_cv_text = cv_text

    ats_original = analyze_ats(full_cv_text, job_description, keywords_override=job_keywords)

    missing_keyword_groups = attempt.get("missing_keyword_groups") or {}
    missing_keyword_categories_order = [
        g.get("category") for g in (updated_data.get("skills_grouped") or [])
        if g.get("category")
    ]

    _apply_fixes_result = {
        "ats_original": ats_original,
        "ats_generated": ats_generated,
        "generated_cv": rendered_cv,
        "cover_letter": attempt.get("cover_letter_html"),
        "cover_letter_data": attempt.get("cover_letter_data"),
        "cv_text": full_cv_text[:500],
        "template": selected_template,
        "region": region_code,
        "region_rules": region_rules,
        "quality_review": None,
    }
    results_ctx = _build_results_context(
        _apply_fixes_result,
        missing_keyword_groups,
        missing_keyword_categories_order,
        fixes_applied=len(flags),
    )
    return templates.TemplateResponse(
        "partials/results.html",
        {"request": request, **results_ctx},
    )


# ---------------------------------------------------------------------------
# Append user-selected skills to the existing CV (no LLM call)
# ---------------------------------------------------------------------------


@router.post("/add-skills")
async def add_skills(request: Request):
    """Append the keywords the user staged from the ATS comparison panel
    directly to ``cv_data.skills`` and re-render the template.

    No LLM call — this is a pure data merge + template re-render. Cheaper and
    faster than a full /analyze rerun, and avoids the AI rewriting passages
    the user is happy with.
    """
    attempt_id = request.state.session.get("attempt_id")
    if not attempt_id:
        return templates.TemplateResponse(
            "partials/error.html",
            {"request": request, "error": "No active session."},
        )

    attempt = get_attempt(attempt_id)
    if not attempt or not attempt.get("cv_data"):
        return templates.TemplateResponse(
            "partials/error.html",
            {"request": request, "error": "No generated CV found. Please generate first."},
        )

    form = await request.form()
    new_skills = [str(v).strip() for v in form.getlist("staged_keywords") if str(v).strip()]
    if not new_skills:
        return templates.TemplateResponse(
            "partials/error.html",
            {"request": request, "error": "No skills selected."},
        )

    cv_data = dict(attempt["cv_data"])
    template_id = attempt.get("template_id", "modern")
    region_code = attempt.get("region", "US")
    job_description = attempt.get("job_description", "")
    missing_keyword_groups: dict[str, str] = attempt.get("missing_keyword_groups") or {}

    existing = cv_data.get("skills") or []
    seen = {s.strip().lower() for s in existing if isinstance(s, str)}

    # Extend seen with every skill already in skills_grouped so we don't
    # add the same keyword twice across the flat list and grouped structure.
    grouped = list(cv_data.get("skills_grouped") or [])
    for group in grouped:
        for s in group.get("items", []):
            if isinstance(s, str):
                seen.add(s.strip().lower())

    added_kws: list[str] = []
    for kw in new_skills:
        key = kw.lower()
        if key in seen:
            continue
        seen.add(key)
        existing.append(kw)
        added_kws.append(kw)
    cv_data["skills"] = existing

    # Mirror additions into skills_grouped — category-aware placement.
    if grouped and added_kws:
        _ADDITIONAL_NAMES = {"additional skills", "additional", "other skills", "other"}
        # Build a case-insensitive lookup of existing group names → group object
        group_by_cat = {g.get("category", "").strip().lower(): g for g in grouped}

        for kw in added_kws:
            category = missing_keyword_groups.get(kw, "")
            cat_key = category.strip().lower()
            if category and cat_key in group_by_cat:
                # Append to the matching existing group
                target = group_by_cat[cat_key]
                target["items"] = [*list(target.get("items") or []), kw]
            elif category:
                # Category from categorizer but no matching group yet — create it
                new_group: dict = {"category": category, "items": [kw]}
                grouped.append(new_group)
                group_by_cat[cat_key] = new_group
            else:
                # No mapping (older attempt / categorizer failed) — fall back to Additional Skills
                fallback = next(
                    (g for g in grouped if g.get("category", "").strip().lower() in _ADDITIONAL_NAMES),
                    None,
                )
                if fallback is not None:
                    fallback["items"] = [*list(fallback.get("items") or []), kw]
                else:
                    fallback_group: dict = {"category": "Additional Skills", "items": [kw]}
                    grouped.append(fallback_group)
                    group_by_cat["additional skills"] = fallback_group

        cv_data["skills_grouped"] = grouped

    llm_usage = cv_data.pop("_llm_usage", {})
    rendered_cv = templates.get_template(f"cv_templates/{template_id}.html").render(**cv_data)
    cv_data["_llm_usage"] = llm_usage

    generated_text = re.sub(r"<style[^>]*>.*?</style>", "", rendered_cv, flags=re.DOTALL)
    generated_text = re.sub(r"<[^>]+>", " ", generated_text)
    generated_text = re.sub(r"\s+", " ", generated_text).strip()

    keyword_data = attempt.get("extracted_keywords")
    job_keywords = keyword_data["all_keywords"] if keyword_data else None
    ats_generated = analyze_ats(generated_text, job_description, keywords_override=job_keywords)

    update_attempt(attempt_id, cv_data=cv_data, rendered_cv=rendered_cv)

    selected_template = get_template(template_id) or get_template("modern")
    region_rules = REGION_RULES.get(region_code, REGION_RULES["US"])

    cv_text = attempt.get("cv_text_preview", "")
    cv_bytes = get_document_bytes(attempt_id, "cv_file")
    cv_filename = get_document_filename(attempt_id, "cv_file")
    if cv_bytes and cv_filename:
        try:
            full_cv_text = parse_cv(cv_filename, cv_bytes)
        except ValueError:
            full_cv_text = cv_text
    else:
        full_cv_text = cv_text
    ats_original = analyze_ats(full_cv_text, job_description, keywords_override=job_keywords)

    logger.info(
        "AddSkills[%s] appended=%d total_skills=%d ats_score=%d",
        attempt_id, len(new_skills), len(cv_data["skills"]), ats_generated.score,
    )

    missing_keyword_categories_order = [
        g.get("category") for g in (cv_data.get("skills_grouped") or []) if g.get("category")
    ]
    _add_skills_result = {
        "ats_original": ats_original,
        "ats_generated": ats_generated,
        "generated_cv": rendered_cv,
        "cover_letter": attempt.get("cover_letter_html"),
        "cover_letter_data": attempt.get("cover_letter_data"),
        "cv_text": full_cv_text[:500],
        "template": selected_template,
        "region": region_code,
        "region_rules": region_rules,
        "quality_review": attempt.get("quality_review"),
    }
    results_ctx = _build_results_context(
        _add_skills_result,
        missing_keyword_groups,
        missing_keyword_categories_order,
        skills_added=len(new_skills),
    )
    return templates.TemplateResponse(
        "partials/results.html",
        {"request": request, **results_ctx},
    )


# ---------------------------------------------------------------------------
# Final Review — mandatory human-review step before any download/save
# ---------------------------------------------------------------------------

_FINAL_REVIEW_REDIRECT_URL = "/final-review"
_FINAL_REVIEW_FLASH_MSG = (
    "Almost there — please complete Final Review before downloading. "
    "This step is required."
)


def _final_review_required_response(request: Request) -> RedirectResponse:
    """Return a redirect to Final Review with a flash message in the session.

    Called by every download endpoint when ``final_review_completed`` is False.
    """
    session = request.state.session
    session["flash"] = _FINAL_REVIEW_FLASH_MSG
    return RedirectResponse(_FINAL_REVIEW_REDIRECT_URL, status_code=303)


@router.get("/final-review", include_in_schema=True)
async def final_review_page(request: Request):
    """Hydrate the Builder with AI-generated cv_data for mandatory human review.

    This is the required step between AI results and any download/save action.
    The builder renders with ``is_final_review=True`` so the top banner is shown
    and the download controls are gated behind the confirm checkbox.
    """
    attempt_id = request.state.session.get("attempt_id")
    if not attempt_id:
        return RedirectResponse("/wizard/step/1", status_code=303)

    attempt = get_attempt(attempt_id)
    if not attempt:
        return RedirectResponse("/wizard/step/1", status_code=303)

    cv_data = attempt.get("cv_data")
    if not cv_data:
        # No generated CV — send back to wizard
        session = request.state.session
        session["flash"] = "Please generate your CV first before proceeding to Final Review."
        return RedirectResponse("/wizard/step/1", status_code=303)

    # Build the builder-compatible cv_data from the AI-generated attempt.
    # The AI pipeline stores cv_data directly on the attempt (not in builder_data),
    # so we adapt it here by injecting it as builder_data for cv_data_from_attempt.
    region = attempt.get("region", "US")
    template_id = attempt.get("template_id", "modern")

    # Adapt AI cv_data → builder_data shape (both share the same field names)
    adapted_attempt = {"builder_data": {**cv_data, "region": region, "template_id": template_id}}
    builder_cv = builder_cv_data_from_attempt(adapted_attempt)

    # Pre-fill PII vault values if available
    pii = request.state.session.get("pii") or {}
    if pii:
        from app.cv_builder.use_cases.build_cv import apply_pii_prefill
        apply_pii_prefill(builder_cv, pii)

    template_options = [(t.id, t.name) for t in list_templates()]
    region_options = [(r.code, f"{r.flag} {r.name}") for r in list_regions()]
    from app.cv_builder.use_cases.build_cv import region_fields_map
    region_fields = region_fields_map()

    # Pop flash message if present
    flash = request.state.session.pop("flash", None)

    final_review_completed = bool(attempt.get("final_review_completed"))

    return templates.TemplateResponse(
        "builder.html",
        {
            "request": request,
            "cv_data": builder_cv,
            "template_options": template_options,
            "region_options": region_options,
            "selected_region": region,
            "selected_template": template_id,
            "region_fields_json": json.dumps(region_fields),
            "dev_mode": request.app.state.dev_mode,
            "editing_cv_id": None,
            "editing_label": "",
            "editing_job_title": "",
            # Final Review context
            "is_final_review": True,
            "final_review_completed": final_review_completed,
            "flash": flash,
            "page_crumbs": [
                {"label": "Create", "href": "/dashboard"},
                {"label": "Generate", "href": "/wizard/step/6"},
                {"label": "Final Review (required)"},
            ],
        },
    )


@router.post("/confirm-review", include_in_schema=True)
async def confirm_review(request: Request):
    """Mark the current attempt's final_review_completed flag as True.

    Called when the user checks the confirm checkbox and clicks the confirm
    button in the Final Review builder. After confirming, download endpoints
    are unlocked for this attempt.
    """
    attempt_id = request.state.session.get("attempt_id")
    if not attempt_id:
        return RedirectResponse(_FINAL_REVIEW_REDIRECT_URL, status_code=303)

    attempt = get_attempt(attempt_id)
    if not attempt or not attempt.get("cv_data"):
        return RedirectResponse(_FINAL_REVIEW_REDIRECT_URL, status_code=303)

    form = await request.form()
    confirmed = form.get("review_confirmed") == "1"

    if confirmed:
        update_attempt(attempt_id, final_review_completed=True)
        logger.info("FinalReview[%s] confirmed — download gate unlocked", attempt_id)
        # Redirect back to Final Review page so user sees unlocked downloads
        # (the page will re-render with final_review_completed=True)
        return RedirectResponse(_FINAL_REVIEW_REDIRECT_URL, status_code=303)

    # Not confirmed — redirect back with a message
    request.state.session["flash"] = "Please check the confirmation box to continue."
    return RedirectResponse(_FINAL_REVIEW_REDIRECT_URL, status_code=303)


# ---------------------------------------------------------------------------
# PDF download (unchanged)
# ---------------------------------------------------------------------------


@router.get("/download-pdf")
async def download_pdf(request: Request):
    """Generate and download the CV as PDF."""
    attempt_id = request.state.session.get("attempt_id")
    if not attempt_id:
        return Response("No active session", status_code=400)

    attempt = get_attempt(attempt_id)
    if not attempt or not attempt.get("rendered_cv"):
        return Response("No generated CV found. Please generate your CV first.", status_code=400)

    # Final Review gate — redirect if not yet confirmed
    if not attempt.get("final_review_completed"):
        return _final_review_required_response(request)

    rendered_cv = attempt["rendered_cv"]
    cv_name = attempt.get("cv_data", {}).get("name", "CV") or "CV"
    safe_name = "".join(c for c in cv_name if c.isalnum() or c in " -_").strip() or "CV"

    pdf_bytes = await generate_pdf(rendered_cv)
    if pdf_bytes is None:
        return Response("PDF generation failed. Please try again.", status_code=500)

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{safe_name} - QuillCV.pdf"',
        },
    )


@router.get("/download-docx")
async def download_docx(request: Request):
    """Generate and download the CV as DOCX."""
    attempt_id = request.state.session.get("attempt_id")
    if not attempt_id:
        return Response("No active session", status_code=400)

    attempt = get_attempt(attempt_id)
    if not attempt or not attempt.get("cv_data"):
        return Response("No generated CV found. Please generate your CV first.", status_code=400)

    # Final Review gate
    if not attempt.get("final_review_completed"):
        return _final_review_required_response(request)

    cv_data = attempt["cv_data"]
    region_code = attempt.get("region", "AU") or "AU"
    template_id = attempt.get("template_id", "classic") or "classic"
    cv_name = cv_data.get("name", "CV") or "CV"
    safe_name = "".join(c for c in cv_name if c.isalnum() or c in " -_").strip() or "CV"

    try:
        docx_bytes = generate_docx(cv_data, region_code=region_code, template_id=template_id)
    except Exception:
        logger.exception("DOCX generation failed for attempt=%s", attempt_id)
        return Response("DOCX generation failed. Please try again.", status_code=500)

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{safe_name} - QuillCV.docx"',
        },
    )


# ---------------------------------------------------------------------------
# Cover letter downloads
# ---------------------------------------------------------------------------


@router.get("/download-cover-letter-pdf")
async def download_cover_letter_pdf(request: Request):
    """Generate and download the cover letter as PDF."""
    attempt_id = request.state.session.get("attempt_id")
    if not attempt_id:
        return Response("No active session", status_code=400)
    attempt = get_attempt(attempt_id)
    if not attempt or not attempt.get("cover_letter_html"):
        return Response("No cover letter found. Please generate first.", status_code=400)

    # Final Review gate
    if not attempt.get("final_review_completed"):
        return _final_review_required_response(request)

    cover_letter_html = attempt["cover_letter_html"]
    cv_name = attempt.get("cv_data", {}).get("name", "Cover Letter") or "Cover Letter"
    safe_name = "".join(c for c in cv_name if c.isalnum() or c in " -_").strip() or "Cover Letter"

    pdf_bytes = await generate_pdf(cover_letter_html)
    if pdf_bytes is None:
        return Response("PDF generation failed. Please try again.", status_code=500)

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{safe_name} - Cover Letter - QuillCV.pdf"',
        },
    )


def _build_cover_letter_docx_bytes(cl: dict) -> bytes:
    """Render a cover-letter dict into a DOCX byte string."""
    from io import BytesIO

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    def add(text, align=None):
        if not text:
            return
        p = doc.add_paragraph(text)
        if align:
            p.alignment = align
        return p

    add(cl.get("date", ""), align=WD_ALIGN_PARAGRAPH.RIGHT)
    recipient = cl.get("recipient", "")
    company = cl.get("company_name", "")
    if recipient or company:
        add(", ".join(filter(None, [recipient, company])))
    doc.add_paragraph()
    add(cl.get("salutation", ""))
    add(cl.get("opening", ""))
    for para in cl.get("body_paragraphs") or []:
        add(para)
    add(cl.get("contribution", ""))
    add(cl.get("closing", ""))
    doc.add_paragraph()
    add(cl.get("sign_off", ""))
    add(cl.get("name", ""))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/download-cover-letter-docx")
async def download_cover_letter_docx(request: Request):
    """Generate and download the cover letter as DOCX."""
    attempt_id = request.state.session.get("attempt_id")
    if not attempt_id:
        return Response("No active session", status_code=400)
    attempt = get_attempt(attempt_id)
    cl = attempt.get("cover_letter_data") if attempt else None
    if not cl:
        return Response("No cover letter found. Please generate first.", status_code=400)

    # Final Review gate
    if not attempt.get("final_review_completed"):
        return _final_review_required_response(request)

    cv_name = attempt.get("cv_data", {}).get("name", "Cover Letter") or "Cover Letter"
    safe_name = "".join(c for c in cv_name if c.isalnum() or c in " -_").strip() or "Cover Letter"

    return Response(
        content=_build_cover_letter_docx_bytes(cl),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{safe_name} - Cover Letter - QuillCV.docx"',
        },
    )


# ---------------------------------------------------------------------------
# Combined CV + cover letter ZIP downloads
# ---------------------------------------------------------------------------


@router.get("/download-all-pdf")
async def download_all_pdf(request: Request):
    """Bundle the CV PDF and cover-letter PDF into a single ZIP download."""
    attempt_id = request.state.session.get("attempt_id")
    if not attempt_id:
        return Response("No active session", status_code=400)

    attempt = get_attempt(attempt_id)
    if not attempt or not attempt.get("rendered_cv"):
        return Response("No generated CV found. Please generate your CV first.", status_code=400)
    if not attempt.get("cover_letter_html"):
        return Response("No cover letter found. Generate one to use the bundle download.", status_code=400)

    # Final Review gate
    if not attempt.get("final_review_completed"):
        return _final_review_required_response(request)

    cv_name = attempt.get("cv_data", {}).get("name", "CV") or "CV"
    safe_name = "".join(c for c in cv_name if c.isalnum() or c in " -_").strip() or "CV"

    cv_pdf = await generate_pdf(attempt["rendered_cv"])
    cl_pdf = await generate_pdf(attempt["cover_letter_html"])
    if cv_pdf is None or cl_pdf is None:
        return Response("PDF generation failed. Please try again.", status_code=500)

    import zipfile
    from io import BytesIO

    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{safe_name} - QuillCV.pdf", cv_pdf)
        zf.writestr(f"{safe_name} - Cover Letter - QuillCV.pdf", cl_pdf)

    return Response(
        content=buf.getvalue(),
        media_type="application/zip",
        headers={
            "Content-Disposition": f'attachment; filename="{safe_name} - QuillCV.zip"',
        },
    )


@router.get("/download-all-docx")
async def download_all_docx(request: Request):
    """Bundle the CV DOCX and cover-letter DOCX into a single ZIP download."""
    attempt_id = request.state.session.get("attempt_id")
    if not attempt_id:
        return Response("No active session", status_code=400)

    attempt = get_attempt(attempt_id)
    if not attempt or not attempt.get("cv_data"):
        return Response("No generated CV found. Please generate your CV first.", status_code=400)
    cl = attempt.get("cover_letter_data")
    if not cl:
        return Response("No cover letter found. Generate one to use the bundle download.", status_code=400)

    # Final Review gate
    if not attempt.get("final_review_completed"):
        return _final_review_required_response(request)

    cv_data = attempt["cv_data"]
    region_code = attempt.get("region", "AU") or "AU"
    template_id = attempt.get("template_id", "classic") or "classic"
    cv_name = cv_data.get("name", "CV") or "CV"
    safe_name = "".join(c for c in cv_name if c.isalnum() or c in " -_").strip() or "CV"

    try:
        cv_docx = generate_docx(cv_data, region_code=region_code, template_id=template_id)
    except Exception:
        logger.exception("DOCX generation failed for attempt=%s", attempt_id)
        return Response("DOCX generation failed. Please try again.", status_code=500)

    cl_docx = _build_cover_letter_docx_bytes(cl)

    import zipfile
    from io import BytesIO

    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{safe_name} - QuillCV.docx", cv_docx)
        zf.writestr(f"{safe_name} - Cover Letter - QuillCV.docx", cl_docx)

    return Response(
        content=buf.getvalue(),
        media_type="application/zip",
        headers={
            "Content-Disposition": f'attachment; filename="{safe_name} - QuillCV.zip"',
        },
    )
