"""Generate ATS-friendly DOCX files from structured CV data.

Uses python-docx to produce clean Word documents that ATS systems can parse
reliably. No tables for layout, no text boxes — just standard headings,
paragraphs, and bullet lists.

Each template ID maps to a distinct visual style that mirrors the HTML preview:
accent colors, heading fonts, and structural variations are applied per template.
"""

import logging
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# Letter: 8.5 x 11 in; A4: 8.27 x 11.69 in
_LETTER_REGIONS = {"US", "CA", "CO", "VE"}

_PAGE_WIDTH_LETTER = Inches(8.5)
_PAGE_HEIGHT_LETTER = Inches(11)
_PAGE_WIDTH_A4 = Emu(int(8.27 * 914400))
_PAGE_HEIGHT_A4 = Emu(int(11.69 * 914400))
_MARGIN = Inches(1)

# Shared size constants
_SIZE_CONTACT = 10
_SIZE_SECTION = 11
_SIZE_JOB_TITLE = 11
_SIZE_BODY = 10.5


# ---------------------------------------------------------------------------
# Per-template style definitions
# ---------------------------------------------------------------------------

TEMPLATE_STYLES: dict[str, dict] = {
    "classic": {
        "accent": RGBColor(0x2C, 0x3E, 0x50),      # dark slate blue
        "heading_font": "Georgia",
        "body_font": "Georgia",
        "name_size": 22,
        "name_color": RGBColor(0x2C, 0x3E, 0x50),
        "title_color": RGBColor(0x55, 0x55, 0x55),
        "section_border": True,
        "contact_align": WD_ALIGN_PARAGRAPH.LEFT,
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "job_title_color": RGBColor(0x1A, 0x1A, 0x1A),  # near-black, bold
    },
    "modern": {
        "accent": RGBColor(0x1A, 0x73, 0xE8),       # Google blue
        "heading_font": "Calibri",
        "body_font": "Calibri",
        "name_size": 24,
        "name_color": RGBColor(0x1A, 0x73, 0xE8),
        "title_color": RGBColor(0x55, 0x55, 0x55),
        "section_border": True,
        "contact_align": WD_ALIGN_PARAGRAPH.LEFT,
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "job_title_color": RGBColor(0x1A, 0x1A, 0x1A),
        "company_color": RGBColor(0x1A, 0x73, 0xE8),  # company in accent
    },
    "tech": {
        "accent": RGBColor(0x0E, 0xA5, 0xE9),       # sky blue
        "heading_font": "Segoe UI",
        "body_font": "Segoe UI",
        "name_size": 20,
        "name_color": RGBColor(0x0C, 0x4A, 0x6E),   # deep navy
        "title_color": RGBColor(0x0E, 0xA5, 0xE9),  # title in accent
        "section_border": True,                      # light border
        "section_border_color": "E0F2FE",            # very light blue
        "contact_align": WD_ALIGN_PARAGRAPH.LEFT,
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "job_title_color": RGBColor(0x1E, 0x1E, 0x1E),
        "company_color": RGBColor(0x0E, 0xA5, 0xE9),
        "show_tech_stack": True,
        "header_left_bar": True,                     # left bar decoration (visual note)
    },
    "minimal": {
        "accent": RGBColor(0x33, 0x33, 0x33),        # near black
        "heading_font": "Arial",
        "body_font": "Arial",
        "name_size": 20,
        "name_color": RGBColor(0x33, 0x33, 0x33),
        "title_color": RGBColor(0x66, 0x66, 0x66),
        "section_border": False,
        "contact_align": WD_ALIGN_PARAGRAPH.CENTER,
        "name_align": WD_ALIGN_PARAGRAPH.CENTER,
        "job_title_color": RGBColor(0x33, 0x33, 0x33),
    },
    "executive": {
        "accent": RGBColor(0x1A, 0x1A, 0x2E),        # dark navy
        "heading_font": "Georgia",
        "body_font": "Georgia",
        "name_size": 24,
        "name_color": RGBColor(0x1A, 0x1A, 0x2E),
        "title_color": RGBColor(0x44, 0x44, 0x44),
        "section_border": True,
        "contact_align": WD_ALIGN_PARAGRAPH.CENTER,
        "name_align": WD_ALIGN_PARAGRAPH.CENTER,
        "job_title_color": RGBColor(0x1A, 0x1A, 0x2E),
    },
    "compact": {
        "accent": RGBColor(0x44, 0x44, 0x44),
        "heading_font": "Arial Narrow",
        "body_font": "Arial Narrow",
        "name_size": 18,
        "name_color": RGBColor(0x44, 0x44, 0x44),
        "title_color": RGBColor(0x66, 0x66, 0x66),
        "section_border": False,
        "body_size": 9.5,
        "contact_align": WD_ALIGN_PARAGRAPH.LEFT,
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "job_title_color": RGBColor(0x44, 0x44, 0x44),
    },
    "academic": {
        "accent": RGBColor(0x4A, 0x23, 0x6D),        # deep purple
        "heading_font": "Times New Roman",
        "body_font": "Times New Roman",
        "name_size": 22,
        "name_color": RGBColor(0x4A, 0x23, 0x6D),
        "title_color": RGBColor(0x55, 0x55, 0x55),
        "section_border": True,
        "contact_align": WD_ALIGN_PARAGRAPH.CENTER,
        "name_align": WD_ALIGN_PARAGRAPH.CENTER,
        "job_title_color": RGBColor(0x1A, 0x1A, 0x1A),
    },
    "healthcare": {
        "accent": RGBColor(0x00, 0x89, 0x7B),        # teal
        "heading_font": "Calibri",
        "body_font": "Calibri",
        "name_size": 22,
        "name_color": RGBColor(0x00, 0x89, 0x7B),
        "title_color": RGBColor(0x55, 0x55, 0x55),
        "section_border": True,
        "contact_align": WD_ALIGN_PARAGRAPH.LEFT,
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "job_title_color": RGBColor(0x1A, 0x1A, 0x1A),
    },
    "legal": {
        "accent": RGBColor(0x1C, 0x2B, 0x4A),        # dark navy
        "heading_font": "Times New Roman",
        "body_font": "Times New Roman",
        "name_size": 22,
        "name_color": RGBColor(0x1C, 0x2B, 0x4A),
        "title_color": RGBColor(0x55, 0x55, 0x55),
        "section_border": True,
        "contact_align": WD_ALIGN_PARAGRAPH.CENTER,
        "name_align": WD_ALIGN_PARAGRAPH.CENTER,
        "job_title_color": RGBColor(0x1C, 0x2B, 0x4A),
    },
    "creative": {
        "accent": RGBColor(0x7C, 0x3A, 0xED),        # purple
        "heading_font": "Calibri",
        "body_font": "Calibri",
        "name_size": 24,
        "name_color": RGBColor(0x7C, 0x3A, 0xED),
        "title_color": RGBColor(0x55, 0x55, 0x55),
        "section_border": True,
        "contact_align": WD_ALIGN_PARAGRAPH.CENTER,
        "name_align": WD_ALIGN_PARAGRAPH.CENTER,
        "job_title_color": RGBColor(0x1A, 0x1A, 0x1A),
    },
    "sales": {
        "accent": RGBColor(0x16, 0xA3, 0x4A),        # green
        "heading_font": "Calibri",
        "body_font": "Calibri",
        "name_size": 22,
        "name_color": RGBColor(0x16, 0xA3, 0x4A),
        "title_color": RGBColor(0x55, 0x55, 0x55),
        "section_border": True,
        "contact_align": WD_ALIGN_PARAGRAPH.LEFT,
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "job_title_color": RGBColor(0x1A, 0x1A, 0x1A),
    },
    "engineering": {
        "accent": RGBColor(0x1E, 0x40, 0xAF),        # dark blue
        "heading_font": "Arial",
        "body_font": "Arial",
        "name_size": 22,
        "name_color": RGBColor(0x1E, 0x40, 0xAF),
        "title_color": RGBColor(0x44, 0x44, 0x44),
        "section_border": True,
        "contact_align": WD_ALIGN_PARAGRAPH.LEFT,
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "job_title_color": RGBColor(0x1A, 0x1A, 0x1A),
    },
    "education": {
        "accent": RGBColor(0xD9, 0x77, 0x06),        # amber
        "heading_font": "Calibri",
        "body_font": "Calibri",
        "name_size": 22,
        "name_color": RGBColor(0xD9, 0x77, 0x06),
        "title_color": RGBColor(0x55, 0x55, 0x55),
        "section_border": True,
        "contact_align": WD_ALIGN_PARAGRAPH.LEFT,
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "job_title_color": RGBColor(0x1A, 0x1A, 0x1A),
    },
    "consulting": {
        "accent": RGBColor(0x1E, 0x29, 0x3C),        # very dark navy
        "heading_font": "Georgia",
        "body_font": "Calibri",
        "name_size": 24,
        "name_color": RGBColor(0xFF, 0xFF, 0xFF),    # white — handled specially
        "title_color": RGBColor(0xCC, 0xCC, 0xCC),
        "section_border": True,
        "contact_align": WD_ALIGN_PARAGRAPH.CENTER,
        "name_align": WD_ALIGN_PARAGRAPH.CENTER,
        "job_title_color": RGBColor(0x1A, 0x1A, 0x1A),
        "dark_header": True,                          # shaded name block
    },
    "nonprofit": {
        "accent": RGBColor(0x92, 0x40, 0x0E),        # warm earth
        "heading_font": "Calibri",
        "body_font": "Calibri",
        "name_size": 22,
        "name_color": RGBColor(0x92, 0x40, 0x0E),
        "title_color": RGBColor(0x55, 0x55, 0x55),
        "section_border": True,
        "contact_align": WD_ALIGN_PARAGRAPH.LEFT,
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "job_title_color": RGBColor(0x1A, 0x1A, 0x1A),
    },
    "federal": {
        "accent": RGBColor(0x1F, 0x2D, 0x3D),        # very dark
        "heading_font": "Times New Roman",
        "body_font": "Times New Roman",
        "name_size": 20,
        "name_color": RGBColor(0x1F, 0x2D, 0x3D),
        "title_color": RGBColor(0x44, 0x44, 0x44),
        "section_border": True,
        "contact_align": WD_ALIGN_PARAGRAPH.LEFT,
        "name_align": WD_ALIGN_PARAGRAPH.LEFT,
        "job_title_color": RGBColor(0x1F, 0x2D, 0x3D),
    },
    "hoja-de-vida": {
        "accent": RGBColor(0xB4, 0x78, 0x18),        # amber/gold
        "heading_font": "Calibri",
        "body_font": "Calibri",
        "name_size": 22,
        "name_color": RGBColor(0xB4, 0x78, 0x18),
        "title_color": RGBColor(0x55, 0x55, 0x55),
        "section_border": True,
        "contact_align": WD_ALIGN_PARAGRAPH.CENTER,
        "name_align": WD_ALIGN_PARAGRAPH.CENTER,
        "job_title_color": RGBColor(0x1A, 0x1A, 0x1A),
    },
}

# Fallback style used for any unknown template ID
_DEFAULT_STYLE = TEMPLATE_STYLES["classic"]


def _get_style(template_id: str) -> dict:
    return TEMPLATE_STYLES.get(template_id, _DEFAULT_STYLE)


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------

def _set_page_size(doc: Document, region_code: str) -> None:
    """Set page dimensions and margins based on region."""
    use_letter = region_code.upper() in _LETTER_REGIONS
    section = doc.sections[0]

    if use_letter:
        section.page_width = _PAGE_WIDTH_LETTER
        section.page_height = _PAGE_HEIGHT_LETTER
    else:
        section.page_width = _PAGE_WIDTH_A4
        section.page_height = _PAGE_HEIGHT_A4

    section.top_margin = _MARGIN
    section.bottom_margin = _MARGIN
    section.left_margin = _MARGIN
    section.right_margin = _MARGIN


# ---------------------------------------------------------------------------
# Low-level run helpers
# ---------------------------------------------------------------------------

def _add_run(
    para,
    text: str,
    font_name: str,
    size: float,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


# ---------------------------------------------------------------------------
# Section heading
# ---------------------------------------------------------------------------

def _add_section_heading(doc: Document, text: str, style: dict) -> None:
    """Add a styled section heading with optional bottom border."""
    accent: RGBColor = style["accent"]
    heading_font: str = style["heading_font"]
    use_border: bool = style.get("section_border", True)
    border_color: str = style.get("section_border_color", None)

    # Derive a hex string from the accent RGBColor for the border
    if border_color is None:
        border_color = f"{accent[0]:02X}{accent[1]:02X}{accent[2]:02X}"

    para = doc.add_paragraph()
    para.style = doc.styles["Heading 1"]
    run = para.add_run(text.upper())
    run.font.name = heading_font
    run.font.size = Pt(_SIZE_SECTION)
    run.font.color.rgb = accent
    run.bold = True

    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)

    if use_border:
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), border_color)
        pBdr.append(bottom)
        pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Header: name, title, contact
# ---------------------------------------------------------------------------

def _add_header(doc: Document, cv_data: dict, style: dict) -> None:
    """Add the name / title / contact block."""
    name = (cv_data.get("name") or "").strip()
    title = (cv_data.get("title") or "").strip()
    email = (cv_data.get("email") or "").strip()
    phone = (cv_data.get("phone") or "").strip()
    location = (cv_data.get("location") or "").strip()
    linkedin = (cv_data.get("linkedin") or "").strip()
    github = (cv_data.get("github") or "").strip()
    portfolio = (cv_data.get("portfolio") or "").strip()

    heading_font = style["heading_font"]
    body_font = style["body_font"]
    name_color: RGBColor = style["name_color"]
    style["title_color"]
    name_size: float = style["name_size"]
    name_align = style.get("name_align", WD_ALIGN_PARAGRAPH.LEFT)
    contact_align = style.get("contact_align", WD_ALIGN_PARAGRAPH.LEFT)

    if name:
        name_para = doc.add_paragraph()
        name_para.alignment = name_align
        name_para.paragraph_format.space_before = Pt(0)
        name_para.paragraph_format.space_after = Pt(2)
        _add_run(name_para, name, heading_font, name_size, bold=True, color=name_color)

    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = name_align
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(3)
        title_color_actual = style.get("title_color", RGBColor(0x55, 0x55, 0x55))
        _add_run(title_para, title, body_font, 12, color=title_color_actual)

    contact_parts = []
    if email:
        contact_parts.append(email)
    if phone:
        contact_parts.append(phone)
    if location:
        contact_parts.append(location)
    if linkedin:
        contact_parts.append(linkedin)
    if github:
        contact_parts.append(github)
    if portfolio:
        contact_parts.append(portfolio)

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = contact_align
        contact_para.paragraph_format.space_before = Pt(0)
        contact_para.paragraph_format.space_after = Pt(6)
        _add_run(
            contact_para,
            "  |  ".join(contact_parts),
            body_font,
            _SIZE_CONTACT,
            color=RGBColor(0x44, 0x44, 0x44),
        )


# ---------------------------------------------------------------------------
# Experience entry
# ---------------------------------------------------------------------------

def _add_job_entry(doc: Document, exp: dict, style: dict) -> None:
    """Add a single experience entry styled to match the template."""
    title = exp.get("title", "")
    company = exp.get("company", "")
    location = exp.get("location", "")
    date = exp.get("date", "")
    bullets = exp.get("bullets", [])
    tech = exp.get("tech", "")

    body_font = style["body_font"]
    heading_font = style["heading_font"]
    body_size = style.get("body_size", _SIZE_BODY)
    job_title_color: RGBColor = style.get("job_title_color", style["accent"])
    company_color: RGBColor = style.get("company_color", RGBColor(0x55, 0x55, 0x55))
    show_tech = style.get("show_tech_stack", False)

    # Job title
    title_para = doc.add_paragraph()
    title_para.style = doc.styles["Heading 2"]
    title_para.paragraph_format.space_before = Pt(6)
    title_para.paragraph_format.space_after = Pt(0)
    _add_run(title_para, title, heading_font, _SIZE_JOB_TITLE, bold=True, color=job_title_color)

    # Company | location | date line
    meta_parts = []
    if company:
        meta_parts.append(company)
    if location:
        meta_parts.append(location)

    if meta_parts or date:
        meta_para = doc.add_paragraph()
        meta_para.paragraph_format.space_before = Pt(0)
        meta_para.paragraph_format.space_after = Pt(2)

        if meta_parts:
            _add_run(
                meta_para,
                " | ".join(meta_parts),
                body_font,
                body_size,
                italic=True,
                color=company_color,
            )
        if date:
            sep = "   " if meta_parts else ""
            _add_run(
                meta_para,
                f"{sep}{date}",
                body_font,
                body_size,
                color=RGBColor(0x66, 0x66, 0x66),
            )

    # Tech stack (shown on tech template)
    if tech and show_tech:
        tech_para = doc.add_paragraph()
        tech_para.paragraph_format.space_before = Pt(1)
        tech_para.paragraph_format.space_after = Pt(1)
        _add_run(tech_para, "Stack: ", body_font, body_size, bold=True)
        _add_run(tech_para, tech, body_font, body_size, italic=True, color=RGBColor(0x44, 0x44, 0x44))

    # Bullet points
    for bullet in bullets:
        if not bullet:
            continue
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(1)
        bp.paragraph_format.left_indent = Inches(0.25)
        _add_run(bp, bullet, body_font, body_size)

    # Tech stack shown below bullets (non-tech templates that have tech field)
    if tech and not show_tech:
        tech_para = doc.add_paragraph()
        tech_para.paragraph_format.space_before = Pt(1)
        tech_para.paragraph_format.space_after = Pt(2)
        _add_run(tech_para, "Stack: ", body_font, body_size, bold=True)
        _add_run(tech_para, tech, body_font, body_size)


# ---------------------------------------------------------------------------
# Skills section
# ---------------------------------------------------------------------------

def _add_skills(doc: Document, cv_data: dict, style: dict) -> None:
    """Add skills section, with grouped or flat format."""
    skills = cv_data.get("skills") or []
    skills_grouped = cv_data.get("skills_grouped") or []

    if not skills and not skills_grouped:
        return

    _add_section_heading(doc, "Skills", style)

    body_font = style["body_font"]
    body_size = style.get("body_size", _SIZE_BODY)
    accent = style["accent"]

    if skills_grouped:
        for group in skills_grouped:
            if not isinstance(group, dict):
                continue
            cat = group.get("category", "")
            # Support both "skills" and "items" keys from different templates
            grp_skills = group.get("skills") or group.get("items") or []
            if isinstance(grp_skills, list):
                skill_str = ", ".join(str(s) for s in grp_skills if s)
            else:
                skill_str = str(grp_skills)
            if not skill_str:
                continue
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(2)
            sp.paragraph_format.space_after = Pt(1)
            if cat:
                _add_run(sp, f"{cat}: ", body_font, body_size, bold=True, color=accent)
            _add_run(sp, skill_str, body_font, body_size)
    elif skills:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(3)
        sp.paragraph_format.space_after = Pt(6)
        _add_run(sp, ", ".join(str(s) for s in skills if s), body_font, body_size)


# ---------------------------------------------------------------------------
# Education
# ---------------------------------------------------------------------------

def _add_education(doc: Document, cv_data: dict, style: dict) -> None:
    education = cv_data.get("education") or []
    if not education:
        return

    _add_section_heading(doc, "Education", style)

    body_font = style["body_font"]
    body_size = style.get("body_size", _SIZE_BODY)
    accent = style["accent"]

    for edu in education:
        if not isinstance(edu, dict):
            continue
        degree = (edu.get("degree") or "").strip()
        institution = (edu.get("institution") or "").strip()
        date = (edu.get("date") or "").strip()
        if not degree:
            continue
        ep = doc.add_paragraph()
        ep.paragraph_format.space_before = Pt(4)
        ep.paragraph_format.space_after = Pt(1)
        _add_run(ep, degree, body_font, body_size, bold=True, color=accent)
        meta = []
        if institution:
            meta.append(institution)
        if date:
            meta.append(date)
        if meta:
            _add_run(ep, "  —  " + " | ".join(meta), body_font, body_size, color=RGBColor(0x55, 0x55, 0x55))


# ---------------------------------------------------------------------------
# Certifications
# ---------------------------------------------------------------------------

def _add_certifications(doc: Document, cv_data: dict, style: dict) -> None:
    certifications = cv_data.get("certifications") or []
    if not certifications:
        return

    _add_section_heading(doc, "Certifications", style)

    body_font = style["body_font"]
    body_size = style.get("body_size", _SIZE_BODY)

    for cert in certifications:
        cert_text = cert.strip() if isinstance(cert, str) else str(cert)
        if not cert_text:
            continue
        cp = doc.add_paragraph(style="List Bullet")
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(1)
        cp.paragraph_format.left_indent = Inches(0.25)
        _add_run(cp, cert_text, body_font, body_size)


# ---------------------------------------------------------------------------
# Projects
# ---------------------------------------------------------------------------

def _add_projects(doc: Document, cv_data: dict, style: dict) -> None:
    projects = cv_data.get("projects") or []
    if not projects:
        return

    _add_section_heading(doc, "Projects", style)

    body_font = style["body_font"]
    body_size = style.get("body_size", _SIZE_BODY)
    accent = style["accent"]

    for proj in projects:
        if not isinstance(proj, dict):
            continue
        proj_name = (proj.get("name") or proj.get("title") or "").strip()
        proj_desc = (proj.get("description") or proj.get("summary") or "").strip()
        proj_tech = proj.get("tech") or proj.get("stack") or ""
        if isinstance(proj_tech, list):
            proj_tech = ", ".join(str(t) for t in proj_tech if t)
        proj_tech = proj_tech.strip() if isinstance(proj_tech, str) else ""
        proj_url = (proj.get("url") or proj.get("link") or "").strip()
        if not proj_name:
            continue
        pp = doc.add_paragraph()
        pp.paragraph_format.space_before = Pt(4)
        pp.paragraph_format.space_after = Pt(1)
        _add_run(pp, proj_name, body_font, body_size, bold=True, color=accent)
        if proj_url:
            _add_run(pp, f"  ({proj_url})", body_font, body_size, color=RGBColor(0x44, 0x44, 0x44))
        if proj_desc:
            dp = doc.add_paragraph()
            dp.paragraph_format.space_before = Pt(0)
            dp.paragraph_format.space_after = Pt(1)
            _add_run(dp, proj_desc, body_font, body_size)
        if proj_tech:
            tp = doc.add_paragraph()
            tp.paragraph_format.space_before = Pt(0)
            tp.paragraph_format.space_after = Pt(2)
            _add_run(tp, "Technologies: ", body_font, body_size, bold=True)
            _add_run(tp, proj_tech, body_font, body_size)


# ---------------------------------------------------------------------------
# Languages
# ---------------------------------------------------------------------------

def _add_languages(doc: Document, cv_data: dict, style: dict) -> None:
    languages = cv_data.get("languages") or []
    if not languages:
        return

    _add_section_heading(doc, "Languages", style)

    body_font = style["body_font"]
    body_size = style.get("body_size", _SIZE_BODY)

    lang_strs = []
    for lang in languages:
        if isinstance(lang, dict):
            lname = lang.get("language") or lang.get("name") or ""
            level = lang.get("level") or lang.get("proficiency") or ""
            lang_strs.append(f"{lname} ({level})" if level else lname)
        elif isinstance(lang, str):
            lang_strs.append(lang)
    if lang_strs:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(3)
        lp.paragraph_format.space_after = Pt(6)
        _add_run(lp, ", ".join(lang for lang in lang_strs if lang), body_font, body_size)


# ---------------------------------------------------------------------------
# References
# ---------------------------------------------------------------------------

def _add_references(doc: Document, cv_data: dict, style: dict) -> None:
    references = cv_data.get("references") or []
    if not references:
        return

    _add_section_heading(doc, "References", style)

    body_font = style["body_font"]
    body_size = style.get("body_size", _SIZE_BODY)
    accent = style["accent"]

    for ref in references:
        if not isinstance(ref, dict):
            continue
        ref_name = (ref.get("name") or "").strip()
        ref_title = (ref.get("title") or "").strip()
        ref_company = (ref.get("company") or "").strip()
        ref_contact = (ref.get("contact") or "").strip()
        if not ref_name:
            continue
        rp = doc.add_paragraph()
        rp.paragraph_format.space_before = Pt(4)
        rp.paragraph_format.space_after = Pt(1)
        _add_run(rp, ref_name, body_font, body_size, bold=True, color=accent)
        meta = []
        if ref_title:
            meta.append(ref_title)
        if ref_company:
            meta.append(ref_company)
        if ref_contact:
            meta.append(ref_contact)
        if meta:
            _add_run(rp, "  —  " + " | ".join(meta), body_font, body_size, color=RGBColor(0x55, 0x55, 0x55))


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_docx(cv_data: dict, region_code: str = "AU", template_id: str = "classic") -> bytes:
    """Generate an ATS-friendly DOCX from structured CV data.

    Applies template-specific styling (accent colors, fonts, borders) to match
    the HTML preview as closely as Word's format allows.

    Returns the .docx file as bytes.
    """
    if not cv_data:
        cv_data = {}

    style = _get_style(template_id)

    doc = Document()
    _set_page_size(doc, region_code)

    # Reset built-in styles to our chosen body font so Word's defaults don't
    # override our explicit run formatting.
    body_font = style["body_font"]
    for style_name in ("Normal", "Heading 1", "Heading 2"):
        doc.styles[style_name].font.name = body_font

    # -------------------------------------------------------------------------
    # Header
    # -------------------------------------------------------------------------
    _add_header(doc, cv_data, style)

    # -------------------------------------------------------------------------
    # Summary
    # -------------------------------------------------------------------------
    summary = (cv_data.get("summary") or "").strip()
    if summary:
        _add_section_heading(doc, "Professional Summary", style)
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(3)
        sp.paragraph_format.space_after = Pt(6)
        _add_run(sp, summary, style["body_font"], style.get("body_size", _SIZE_BODY))

    # -------------------------------------------------------------------------
    # Skills (before experience for tech template — mirrors HTML)
    # -------------------------------------------------------------------------
    if template_id in ("tech",):
        _add_skills(doc, cv_data, style)

    # -------------------------------------------------------------------------
    # Experience
    # -------------------------------------------------------------------------
    experience = cv_data.get("experience") or []
    if experience:
        _add_section_heading(doc, "Experience", style)
        for exp in experience:
            if isinstance(exp, dict):
                _add_job_entry(doc, exp, style)

    # -------------------------------------------------------------------------
    # Skills (all other templates — after experience)
    # -------------------------------------------------------------------------
    if template_id not in ("tech",):
        _add_skills(doc, cv_data, style)

    # -------------------------------------------------------------------------
    # Education
    # -------------------------------------------------------------------------
    _add_education(doc, cv_data, style)

    # -------------------------------------------------------------------------
    # Certifications
    # -------------------------------------------------------------------------
    _add_certifications(doc, cv_data, style)

    # -------------------------------------------------------------------------
    # Projects
    # -------------------------------------------------------------------------
    _add_projects(doc, cv_data, style)

    # -------------------------------------------------------------------------
    # Languages
    # -------------------------------------------------------------------------
    _add_languages(doc, cv_data, style)

    # -------------------------------------------------------------------------
    # References
    # -------------------------------------------------------------------------
    _add_references(doc, cv_data, style)

    # -------------------------------------------------------------------------
    # Serialise to bytes
    # -------------------------------------------------------------------------
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
